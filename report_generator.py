from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import os
from datetime import datetime

REPORT_DIR = "reports"
os.makedirs(REPORT_DIR, exist_ok=True)


def set_background(document, color="E7F3FF"):
    section = document.sections[0]
    sectPr = section._sectPr
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), color)
    sectPr.append(bg)

def create_graph():
    marks = [70, 80, 65, 90]
    subjects = ["Maths", "CS", "Physics", "English"]

    plt.figure()
    plt.bar(subjects, marks)
    plt.title("Performance Graph")
    plt.savefig("graph.png")
    plt.close()


def create_report():
    name = input("Student Name: ")
    enroll = input("Enrollment No: ")
    topic = input("Report Topic: ")

    doc = Document()
    set_background(doc)

    title = doc.add_heading(topic, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

   
    info = doc.add_paragraph()
    info.add_run(f"Name: ").bold = True
    info.add_run(name + "\n")
    info.add_run("Enrollment: ").bold = True
    info.add_run(enroll + "\n")
    info.add_run("Date: ").bold = True
    info.add_run(datetime.now().strftime("%d-%m-%Y"))

  
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This report explains the topic in a structured and academic manner.")

    doc.add_heading("Objectives", level=2)
    doc.add_paragraph("• Understand the topic\n• Learn practical implementation\n• Analyze results")

   
    doc.add_heading("Data Table", level=2)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "S.No"
    table.rows[0].cells[1].text = "Parameter"
    table.rows[0].cells[2].text = "Value"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Sample"
    table.rows[1].cells[2].text = "Demo"

   
    create_graph()
    doc.add_heading("Graph Analysis", level=2)
    doc.add_picture("graph.png", width=Inches(4))

   
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{name} | {enroll}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filename = os.path.join(REPORT_DIR, topic.replace(" ", "_") + ".docx")
    doc.save(filename)
    os.remove("graph.png")

    print("✅ Report Created Successfully")


def update_report():
    file = input("Enter report name to update (without .docx): ")
    path = os.path.join(REPORT_DIR, file + ".docx")

    if not os.path.exists(path):
        print("❌ Report not found")
        return

    doc = Document(path)
    doc.add_heading("Updated Section", level=2)
    doc.add_paragraph("This section was added during update.")
    doc.save(path)

    print("✅ Report Updated")


def delete_report():
    file = input("Enter report name to delete (without .docx): ")
    path = os.path.join(REPORT_DIR, file + ".docx")

    if os.path.exists(path):
        os.remove(path)
        print("🗑 Report Deleted")
    else:
        print("❌ File not found")


def menu():
    while True:
        print("""
------ DOC REPORT GENERATOR ------
1. Create Report
2. Update Report
3. Delete Report
4. Exit
---------------------------------
""")
        choice = input("Enter choice: ")

        if choice == "1":
            create_report()
        elif choice == "2":
            update_report()
        elif choice == "3":
            delete_report()
        elif choice == "4":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice")

menu()
